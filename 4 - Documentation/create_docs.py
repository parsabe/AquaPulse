import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, text, title="NOTE / SPECIAL SECTION", border_color="008080", bg_color="F0F8FF"):
    """Adds a stylish callout box for notes or custom sections."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0, 51, 102)
    
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10.5)
    run_b.font.color.rgb = RGBColor(51, 51, 51)
    
    # Empty line after box
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def build_word_document(output_path):
    doc = Document()
    
    # Page setup - Margins 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles configuration
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(40, 40, 40)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("AquaPulse Underwater Vision & Telemetry System")
    r_title.bold = True
    r_title.font.name = "Calibri Light"
    r_title.font.size = Pt(26)
    r_title.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("Comprehensive Technical Report: Phase 0 (Preprocessing Pipeline) & Phase 3 (AI Process & Telemetry Engine)")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    # Metadata Box
    tbl_meta = doc.add_table(rows=2, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Project:", " AquaPulse Spreewald Ecosystem"), ("Document Type:", " Technical Architecture & Workflow Specification")],
        [("Target Directories:", " 0 - Preprocessing & 3 - AI process"), ("Status:", " Completed & Synchronized")]
    ]
    for r_idx, row_content in enumerate(meta_data):
        for c_idx, (label, val) in enumerate(row_content):
            cell = tbl_meta.cell(r_idx, c_idx)
            cell.width = Inches(3.25)
            set_cell_background(cell, "F5F7FA")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r_l = p.add_run(label)
            r_l.bold = True
            r_l.font.size = Pt(9.5)
            r_l.font.color.rgb = RGBColor(0, 51, 102)
            r_v = p.add_run(val)
            r_v.font.size = Pt(9.5)
            r_v.font.color.rgb = RGBColor(60, 60, 60)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Helper for Headings
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Calibri Light"
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0, 51, 102)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0, 102, 153)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(51, 51, 51)
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0, 51, 102)
        r2 = p.add_run(text)
        return p

    # --- SECTION 1: EXECUTIVE OVERVIEW ---
    add_h1("1. Executive Overview & System Architecture")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "The AquaPulse project represents an end-to-end computer vision and artificial intelligence framework "
        "engineered specifically for non-invasive marine life monitoring, fish species classification, multi-object tracking, "
        "and automated ecological telemetry in underwater environments (such as the Spreewald river network). "
        "Raw underwater camera captures suffer from heavy optical distortions, illumination attenuation, backscattering, "
        "and background clutter. To solve these challenges, AquaPulse implements a modular two-tier architecture: "
        "a GPU-accelerated Preprocessing Pipeline (Phase 0) followed by a Deep Learning AI Inference Engine (Phase 3)."
    )
    
    add_bullet("Phase 0 (Preprocessing): ", "Transforms raw dual-fisheye 360° videos into standard rectilinear MP4 streams, extracts sequential frames, applies YOLOv8 filtering to isolate frames containing fish, normalizes spatial dimensions to 512x512 to prevent VRAM Out-Of-Memory (OOM) errors, and applies Multi-Scale Retinex with Color Restoration (MSRCR) underwater dehazing.")
    add_bullet("Phase 3 (AI Process): ", "Integrates YOLOv8 and Vision Transformers (ViT) with Bayesian Neural Networks (BNN) and Domain Adversarial Neural Networks (DANN) for domain adaptation. Incorporates ByteTrack multi-object tracking, GBIF and Wikimedia live taxonomic API retrieval, local LLM telemetry via Ollama (Llama 3), and multi-lingual voice synthesis with a Cyberpunk HUD interface.")

    # --- SECTION 2: PHASE 0 - PREPROCESSING PIPELINE ---
    add_h1("2. Phase 0: Preprocessing Pipeline Breakdown")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "The preprocessing stage converts unrefined high-resolution dual-fisheye recordings into high-quality, "
        "optically enhanced, spatially uniform image datasets ready for deep learning model training."
    )
    
    add_h2("2.1. Dual-Fisheye to Rectilinear MP4 Conversion (insv_mp4.py)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "Raw 360° footage captured by Insta360 cameras is stored in dual-fisheye format (.insv files), containing severe spherical distortion "
        "and circular image masks. To make these videos usable for standard vision models, insv_mp4.py executes a CUDA PyTorch geometric transformation:"
    )
    add_bullet("Mathematical Mapping: ", "The left circle center (cx = W_in/4, cy = H_in/2) and radius (R = H_in/2) are established. Focal length f = (W_out / 2) / tan(FOV_x / 2) calculates 3D rectilinear rays (x, y). Equidistant projection computes r_fish = f_fish * theta, mapping output pixels back to input fisheye coordinates.")
    add_bullet("CUDA Grid Sampling: ", "Normalized coordinates in range [-1, 1] are passed to torch.nn.functional.grid_sample(..., mode='bilinear', align_corners=True). This completes frame remapping natively on GPU VRAM, outputting a clear, flattened 1920x1080 30FPS MP4 video.")

    add_h2("2.2. Frame Extraction Process (1 - all_frames.py)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "Extracted high-resolution videos are processed sequentially using OpenCV VideoCapture. Frames are extracted across video files, "
        "formatted with relative directory prefixes, and saved as zero-padded sequential image files (e.g., <prefix>_frame_000123.jpg) into the 'All Frames' directory."
    )

    add_h2("2.3. YOLO-Based Fish Frame Filtering (2 - only_fish_frames.py)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "Raw underwater recordings contain extended periods where no fish are present (empty water, plant background, sand bed). "
        "Including empty frames in training datasets introduces background noise and wastes compute resources."
    )
    add_bullet("Automated Filtering: ", "2 - only_fish_frames.py loads a CUDA-accelerated YOLOv8 model (yolov8n.pt) and evaluates extracted frames in batches of 32.")
    add_bullet("Target Verification: ", "Only frames with bounding box confidence detections corresponding to target aquatic organisms (fish class) are copied into 'Extracted fish objects in frames'. All empty frames are automatically excluded from downstream training.")

    add_h2("2.4. Spatial Normalization & Resizing to 512x512 (3 - resize.py)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "High-resolution video frames (1080p or 4K) place enormous memory demands on GPU VRAM during multi-epoch neural network training "
        "and vision transformer attention computation, frequently triggering Out-Of-Memory (OOM) fatal errors."
    )
    add_bullet("Memory & Resolution Optimization: ", "To guarantee stable training across modern GPU architectures, 3 - resize.py converts all frames into PyTorch CUDA float32 tensors and executes bilinear tensor interpolation (torch.nn.functional.interpolate), resizing all images uniformly to 512 x 512 pixels.")
    add_bullet("OOM Prevention: ", "Standardizing all inputs to 512x512 limits VRAM allocation per batch while retaining fine morphological features (scales, fins, body shape) essential for species classification.")

    add_h2("2.5. Underwater Image & Video Dehazing — MSRCR Method")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "Underwater optical environments suffer from severe wavelength-dependent absorption (red wavelengths decay within 5 meters), "
        "forward scattering (blurring object edges), and backscattering (reducing image contrast). To restore visibility, AquaPulse implements "
        "Multi-Scale Retinex with Color Restoration (MSRCR):"
    )
    add_bullet("Retinex Model: ", "Decomposes image I(x,y) into illumination L(x,y) and surface reflectance R(x,y): I(x,y) = R(x,y) * L(x,y). Logarithmic transformation converts multiplication into addition: log R = log I - log L.")
    add_bullet("Multi-Scale Gaussian Surround: ", "Computes Gaussian blurs across three distinct scale sigmas [15, 80, 250] to capture fine details, medium contrast, and global illumination: R_MSR = sum( log I - log(F_k * I) ) / 3.")
    add_bullet("Color Restoration Factor (CRF): ", "Multiplies Retinex output by a color adjustment weight C_i(x,y) = log(125 * I_i) - log(sum(I_channels)), removing green/blue color cast and restoring natural underwater hues.")
    add_bullet("CUDA Acceleration: ", "Implemented via PyTorch separable Gaussian convolutions (gaussian_blur_gpu) directly on GPU tensors, executing batch video dehazing at 30+ FPS while preserving video audio tracks via FFmpeg.")

    add_callout_box(
        doc,
        "This section is reserved for custom user logs, specific underwater camera calibration parameters, "
        "water salinity/turbidity notes, or fine-tuned MSRCR hyperparameter overrides (e.g., custom scale sigmas or gain/offset values).",
        title="EMPTY SECTION FOR DEHAZING & COLOR RESTORATION OVERRIDES",
        border_color="008080",
        bg_color="F9FBFD"
    )

    # --- SECTION 3: PHASE 3 - AI PROCESS ---
    add_h1("3. Phase 3: AI Process & Neural Telemetry Engine")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "The AI Process directory (3 - AI process) contains the complete real-time tracking, species classification, "
        "biomedical knowledge retrieval, and LLM telemetry synthesis engine."
    )

    add_h2("3.1. Cyberpunk HUD Telemetry Interface (main.py)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "main.py implements a real-time computer vision engine wrapped in a futuristic Cyberpunk HUD overlay. "
        "It features live hardware detection across NVIDIA CUDA, Apple Metal MPS, AMD DirectML, and CPU engines. "
        "The interface renders animated scanlines, dynamic corner brackets, confidence telemetry bars, and interactive target locks."
    )

    add_h2("3.2. Model Backbones & Domain Adaptation (ViT, BNN, DANN)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "The AI engine combines multiple neural architectures to achieve robust identification under varying water clarity conditions:"
    )
    add_bullet("YOLOv8 & Vision Transformers (ViT): ", "YOLOv8 provides high-speed spatial bounding box localization and initial classification. Vision Transformers split fish images into 16x16 visual patches, utilizing self-attention mechanisms to capture complex texture patterns regardless of occlusion.")
    add_bullet("Bayesian Neural Networks (BNN): ", "Replaces static linear weights with Gaussian distributions (mu, rho) using the reparameterization trick. Computes Evidence Lower Bound (ELBO) loss and KL divergence, providing epistemic uncertainty quantification for rare or out-of-distribution species.")
    add_bullet("Domain Adversarial Neural Networks (DANN): ", "Incorporates a Gradient Reversal Layer (GRL) between the YOLO feature extractor and domain classifier. During training, GRL reverses gradients (multiplies by -alpha), forcing the feature extractor to learn domain-invariant representations that bridge clear tank training data and murky Spreewald river environments.")

    add_h2("3.3. Multi-Object Tracking & Target Lock (ByteTrack / BoTSORT)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "To maintain persistent target identities across video sequences, main.py integrates multi-object tracking logic. "
        "It maintains motion state vectors via Kalman filtering and Hungarian matching algorithms. A sticky target lock system "
        "(update_sticky_target_lock) tracks selected targets using Euclidean distance matching with a lost-frame buffer, "
        "preventing identity swaps during brief fish occlusions."
    )

    add_h2("3.4. Biomedical & Taxonomic Knowledge Retrieval (GBIF & Wikimedia)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "When a fish target is locked, main.py queries the Global Biodiversity Information Facility (GBIF) REST API "
        "(https://api.gbif.org/v1/species/match) to fetch live scientific taxonomy (Kingdom, Phylum, Class, Order, Family, Genus). "
        "High-resolution reference images are retrieved from GBIF occurrences and Wikimedia Commons, then saved to a disk-backed memory cache "
        "(species_snapshots/wiki_cache/) to allow seamless offline operation."
    )

    add_h2("3.5. Local LLM Telemetry Agent (Ollama LLM - Llama 3)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "AquaPulse integrates a local LLM telemetry agent powered by Ollama (Llama 3). "
        "The agent assumes the persona of world-renowned marine biologist Dr. Daniel Pauly. "
        "Upon target lock or user query, the LLM processes live GBIF taxonomy and generates concise, scientifically authoritative "
        "telemetry explanations capped strictly at 40 words."
    )

    add_h2("3.6. Voice Synthesis & GUI Deployment")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "Voice synthesis is handled by a CancellableAudioEngine (pyttsx3) supporting instant voice cancellation and dual-language mode (English/German). "
        "A mutual exclusion lock prevents audio overlap. The application offers a native OS drag-and-drop file selector (tkinter) "
        "and is packaged as a standalone executable via PyInstaller (AquaPulseVision.spec)."
    )

    # --- SECTION 4: WORKFLOW & DIRECTORY MAP ---
    add_h1("4. Complete Workflow & Directory Structure")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("The complete system hierarchy across the workspace is structured as follows:")

    dir_tree = (
        "C:\\Users\\parsa\\Desktop\\Code\\\n"
        "├── 0 - Preprocessing/\n"
        "│   ├── 0 - Source Codes/\n"
        "│   │   ├── Extract Frames/          # 1 - all_frames.py, 2 - only_fish_frames.py, 3 - resize.py\n"
        "│   │   ├── Image dehazing.../      # dehaze-img.py (CUDA MSRCR for single images)\n"
        "│   │   └── Video dehazing.../      # batch_msrcr_dehaze.py (CUDA MSRCR batch video engine)\n"
        "│   └── 1 - Insta 360 Process/      # insv_mp4.py (Dual-fisheye to 1080p MP4 conversion)\n"
        "├── 1 - Training/                   # train_yolo.py, Parsa - yolo + DAT+ DANN + BNN.py\n"
        "├── 2 - Evaluation/                 # Evaluation scripts & validation metrics\n"
        "├── 3 - AI process/\n"
        "│   └── Vision transfromers/        # main.py (Cyberpunk HUD, ViT, MOT, GBIF, Ollama, TTS)\n"
        "└── 4 - Documentation/              # Generated reports, diagrams & documentation\n"
    )
    p_code = doc.add_paragraph()
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(10)
    r_c = p_code.add_run(dir_tree)
    r_c.font.name = "Consolas"
    r_c.font.size = Pt(9.5)
    r_c.font.color.rgb = RGBColor(0, 51, 102)

    # Workflow Summary Table
    add_h2("4.1. System Pipeline Summary Table")
    tbl_flow = doc.add_table(rows=7, cols=4)
    tbl_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Stage", "Script / Module", "Core Technology", "Primary Purpose"]
    for c_idx, text in enumerate(headers):
        cell = tbl_flow.cell(0, c_idx)
        set_cell_background(cell, "003366")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)
        
    flow_data = [
        ("0.1 INSV Conversion", "insv_mp4.py", "PyTorch grid_sample (CUDA)", "Dual-fisheye 360° to rectilinear 1080p MP4"),
        ("0.2 Frame Extract", "1 - all_frames.py", "OpenCV VideoCapture", "Recursive sequential frame sampling"),
        ("0.3 Fish Filtering", "2 - only_fish_frames.py", "YOLOv8 Object Detection", "Isolate frames with fish; strip empty frames"),
        ("0.4 Spatial Resize", "3 - resize.py", "PyTorch Bilinear Tensor Interpolation", "Normalize to 512x512; prevent GPU OOM"),
        ("0.5 MSRCR Dehazing", "batch_msrcr_dehaze.py", "Multi-Scale Retinex (CUDA)", "Underwater haze removal & color upgrade"),
        ("3.0 AI Engine", "main.py", "ViT + BNN + DANN + ByteTrack + Ollama", "Real-time tracking, taxonomy & voice HUD")
    ]
    
    for r_idx, row_items in enumerate(flow_data, start=1):
        for c_idx, text in enumerate(row_items):
            cell = tbl_flow.cell(r_idx, c_idx)
            set_cell_background(cell, "F9FBFD" if r_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECTION 5: REFERENCES ---
    add_h1("5. Comprehensive Scientific References & Literature")
    refs = [
        ("Jobson, D. J., Rahman, Z. U., & Woodell, G. A. (1997). ", "A multiscale retinex for bridging the gap between color automatic images and the visual appearance of scenes. IEEE Transactions on Image Processing, 6(7), 965-976."),
        ("Land, E. H., & McCann, J. J. (1971). ", "Lightness and Retinex theory. Journal of the Optical Society of America, 61(1), 1-11."),
        ("Jocher, G., Chaurasia, A., & Qiu, J. (2023). ", "Ultralytics YOLOv8 (Version 8.0.0) [Computer software]. GitHub. https://github.com/ultralytics/ultralytics"),
        ("Dosovitskiy, A., et al. (2020). ", "An image is worth 16x16 words: Transformers for image recognition at scale. arXiv preprint arXiv:2010.11929."),
        ("Ganin, Y., et al. (2016). ", "Domain-adversarial training of neural networks. Journal of Machine Learning Research, 17(59), 1-35."),
        ("Blundell, C., Cornebise, J., Kavukcuoglu, K., & Wierstra, D. (2015). ", "Weight uncertainty in neural network. International Conference on Machine Learning (ICML), 1613-1622."),
        ("Zhang, Y., et al. (2022). ", "ByteTrack: Multi-object tracking by associating every detection box. European Conference on Computer Vision (ECCV), 1-21."),
        ("Ancuti, C. O., Ancuti, C., Haber, T., & Bekaert, P. (2012). ", "Enhancing underwater images and videos by fusion. IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 81-88."),
        ("GBIF Secretariat. (2026). ", "GBIF Backbone Taxonomy. Checklist dataset https://doi.org/10.15468/39omei accessed via GBIF REST API."),
        ("Touvron, H., et al. (2023). ", "Llama 2 / Llama 3: Open foundation and fine-tuned chat models. Meta AI Research.")
    ]
    
    for authors, title_pub in refs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        r_a = p.add_run(authors)
        r_a.bold = True
        r_a.font.size = Pt(10)
        r_a.font.color.rgb = RGBColor(0, 51, 102)
        r_t = p.add_run(title_pub)
        r_t.font.size = Pt(10)
        r_t.font.italic = True
        r_t.font.color.rgb = RGBColor(60, 60, 60)

    # Save document
    doc.save(output_path)
    print(f"[SUCCESS] Successfully created Word report at: {output_path}")

if __name__ == "__main__":
    out_dir = r"C:\Users\parsa\Desktop\Code\4 - Documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "AquaPulse_Preprocessing_and_AI_Process_Report.docx")
    build_word_document(out_path)
